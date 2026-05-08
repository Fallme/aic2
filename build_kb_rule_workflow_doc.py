from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = "E:/CODE/copenc/saascontract/contract_kb_rule_workflow.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_table(table, header_fill="1F4E79"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = OxmlElement("w:tcMar")
            for side in ["top", "left", "bottom", "right"]:
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:w"), "100")
                node.set(qn("w:type"), "dxa")
                tc_mar.append(node)
            tc_pr.append(tc_mar)
            if row_index == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.color.rgb = RGBColor(255, 255, 255)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.color.rgb = RGBColor(31, 78, 121)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.2
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r2 = p.add_run(text[len(bold_prefix):])
        runs = [r1, r2]
    else:
        runs = [p.add_run(text)]
    for r in runs:
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(10.5)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(10)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color="FFFFFF")
        if widths:
            table.rows[0].cells[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            if widths:
                cells[i].width = Cm(widths[i])
    style_table(table)
    doc.add_paragraph()
    return table


def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(8.5)
    doc.add_paragraph()


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)

styles = doc.styles
styles["Normal"].font.name = "Microsoft YaHei"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
styles["Normal"].font.size = Pt(10.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("合同生成项目工作流设计")
run.bold = True
run.font.name = "Microsoft YaHei"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(31, 78, 121)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("知识库与规则库构建部分")
r.font.name = "Microsoft YaHei"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(89, 89, 89)

meta = doc.add_table(rows=4, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ("适用项目", "AI 合同生成、合同审查、企业制度规则化"),
    ("核心目标", "让用户上传的法规、制度、习惯做法、模板和历史合同沉淀为可检索知识与可执行规则"),
    ("设计原则", "原文可追溯、AI 候选、人工生效、版本可控、调用可审计"),
    ("版本日期", "V1.0 / 2026-04-29"),
]
for i, (k, v) in enumerate(meta_data):
    set_cell_text(meta.rows[i].cells[0], k, bold=True, color="FFFFFF")
    set_cell_shading(meta.rows[i].cells[0], "1F4E79")
    set_cell_text(meta.rows[i].cells[1], v)
style_table(meta, header_fill="1F4E79")
doc.add_page_break()

add_heading(doc, "1. 业务目标与范围", 1)
add_body(doc, "本工作流用于合同生成系统中的知识库与规则库构建环节。系统允许用户上传法律法规、公司规章制度、合同模板、历史合同、审批规范、交易习惯说明、谈判纪要等资料，由 AI 完成解析、抽取、归纳、结构化和候选规则生成，再经过人工审核后沉淀为可在合同生成和合同审查中调用的知识资产。")
add_bullets(doc, [
    "知识库负责保存原文、切片、语义索引、摘要、条款画像和引用定位，解决“依据在哪里”的问题。",
    "规则库负责保存经过结构化和审核的生成约束、审查规则、审批规则和风险规则，解决“合同应当如何生成和判断”的问题。",
    "AI 的角色是辅助提取与生成候选规则，不直接让规则自动生效；法律、审批、金额阈值、权责边界类规则必须由法务或管理员审核。",
])

add_heading(doc, "2. 总体工作流", 1)
add_body(doc, "推荐将该能力设计为一个异步的 Document-to-Knowledge-to-Rule Workflow。用户上传后先进入文档处理队列，系统完成解析和抽取，生成知识库条目与候选规则；候选规则进入审核队列，审核通过后进入可调用规则库。")
add_table(doc, ["阶段", "输入", "AI/系统动作", "输出", "是否需要人工"], [
    ("1. 上传接入", "PDF、Word、Excel、图片、网页、文本", "格式校验、权限校验、去重、生成文档 ID", "原始文档记录", "否"),
    ("2. 文档解析", "原始文件", "OCR、段落识别、表格识别、标题层级识别、页码定位", "结构化文本块", "异常文件需处理"),
    ("3. 文档分类", "结构化文本块", "识别法规、制度、模板、历史合同、习惯说明等类型", "文档类型、业务域、效力等级", "可人工修正"),
    ("4. 知识入库", "文本块、表格、摘要", "切片、向量化、关键词索引、章节索引、摘要生成", "知识库条目与索引", "否"),
    ("5. 规则抽取", "知识切片", "抽取义务、禁止、审批、阈值、必备条款、风险点", "候选规则项", "否"),
    ("6. 规则标准化", "候选规则项", "转为统一 DSL/JSON，绑定触发条件和动作", "标准规则草稿", "否"),
    ("7. 冲突检测", "新规则与现有规则", "检测重复、冲突、覆盖范围、优先级问题", "冲突报告", "高风险需人工"),
    ("8. 人工审核", "规则草稿、原文依据、冲突报告", "法务/管理员确认、修改、驳回、启用", "生效规则", "是"),
    ("9. 调用发布", "生效规则", "建立规则索引，发布到生成和审查服务", "可调用规则包", "否"),
], widths=[2.3, 3.0, 4.0, 3.2, 2.4])

add_heading(doc, "3. 用户上传与资料管理", 1)
add_body(doc, "上传入口需要面向非技术用户设计，重点不是“上传文件”，而是让用户明确资料的业务含义、适用范围和可信等级。")
add_table(doc, ["上传字段", "说明", "示例"], [
    ("资料类型", "用户选择或 AI 自动识别", "法规、公司制度、合同模板、历史合同、交易习惯"),
    ("业务领域", "用于后续规则适用范围判断", "采购、销售、人事、技术服务、租赁"),
    ("适用合同类型", "允许多选", "采购合同、服务合同、劳动合同"),
    ("效力等级", "决定规则优先级", "法律法规 > 公司制度 > 模板 > 历史合同 > 习惯"),
    ("生效日期", "支持制度版本治理", "2026-05-01"),
    ("保密级别", "控制知识召回和权限", "公开、内部、敏感、仅法务可见"),
    ("是否参与规则抽取", "用户可只入知识库不生成规则", "是/否"),
], widths=[3.0, 6.2, 5.7])

add_heading(doc, "4. 知识库构建流程", 1)
add_body(doc, "知识库的目标是形成可靠的原文依据层。所有后续规则、合同条款建议和审查结论都必须能回溯到知识库中的原始片段。")
add_table(doc, ["节点", "处理逻辑", "关键输出"], [
    ("解析节点", "对 Word、PDF、图片、Excel 分别采用文本解析、OCR、表格结构化和附件抽取", "paragraph、table、image_text、page_no"),
    ("结构还原节点", "识别标题树、条款编号、附件、页眉页脚、签署页和表格标题", "title_path、clause_no、block_type"),
    ("混合切片节点", "按标题层级、条款编号、语义完整性和表格行切片，避免固定长度打断条款", "chunk_id、chunk_text、parent_path"),
    ("摘要节点", "对文档、章节、条款生成不同粒度摘要", "doc_summary、section_summary、clause_summary"),
    ("知识标注节点", "打标签并标注业务域、合同类型、规则候选程度、风险关键词", "tags、domain、contract_type"),
    ("索引节点", "建立向量索引、关键词索引、结构索引和引用索引", "embedding_id、keyword_index、citation_index"),
], widths=[3.0, 7.7, 4.2])

add_body(doc, "知识库最小字段建议：", bold_prefix="知识库")
add_code_block(doc, """{
  "chunk_id": "CHUNK_20260429_0001",
  "doc_id": "DOC_PURCHASE_POLICY_001",
  "doc_type": "company_policy",
  "business_domain": "采购",
  "contract_type": ["采购合同", "服务合同"],
  "title_path": "采购管理制度/第三章/审批权限",
  "page_no": 8,
  "chunk_text": "采购金额超过50万元的，应提交总经理办公会审批。",
  "summary": "采购金额超过50万元触发总经理办公会审批。",
  "tags": ["金额阈值", "审批", "采购"],
  "embedding_id": "emb_abc123",
  "permission_level": "internal"
}""")

add_heading(doc, "5. 规则库构建流程", 1)
add_body(doc, "规则库不是文档摘要库。只有能够在合同生成、审查、追问或审批中产生明确动作的内容才进入规则库。规则必须包含触发条件、执行动作、适用范围、优先级、依据来源和审核状态。")
add_table(doc, ["规则类型", "从哪些资料抽取", "在合同生成/审查中的用途"], [
    ("必备条款规则", "法规、模板、制度", "要求合同必须包含付款、验收、违约、保密等条款"),
    ("禁止条款规则", "法规、制度、风险案例", "拦截不允许出现的承诺、免责、付款方式或责任限制"),
    ("审批规则", "公司制度、授权矩阵", "根据金额、主体、期限、风险等级触发审批提醒"),
    ("信息追问规则", "模板、制度、历史合同", "生成前发现缺少金额、交付周期、验收标准时自动追问"),
    ("条款推荐规则", "模板、历史合同、最佳实践", "推荐标准条款或相似场景条款"),
    ("风险提示规则", "审查标准、历史争议、法规", "识别合同草稿中的缺失、冲突、异常或高风险表达"),
    ("生成约束规则", "制度、模板、业务习惯", "限制 AI 不得自由改变付款比例、责任上限、管辖地等关键内容"),
], widths=[3.2, 4.3, 7.4])

add_body(doc, "规则库最小字段建议：", bold_prefix="规则库")
add_code_block(doc, """{
  "rule_id": "RULE_PURCHASE_APPROVAL_0001",
  "rule_name": "采购金额超过50万元需总经理办公会审批",
  "rule_type": "approval_rule",
  "contract_type": ["采购合同", "服务合同"],
  "business_domain": "采购",
  "trigger_condition": {
    "field": "contract_amount",
    "operator": ">",
    "value": 500000,
    "currency": "CNY"
  },
  "action": {
    "type": "require_approval",
    "approval_node": "总经理办公会",
    "message": "本合同金额超过50万元，需提交总经理办公会审批。"
  },
  "priority": 80,
  "risk_level": "high",
  "source": {
    "doc_id": "DOC_PURCHASE_POLICY_001",
    "chunk_id": "CHUNK_20260429_0001",
    "page_no": 8,
    "title_path": "采购管理制度/第三章/审批权限",
    "quote": "采购金额超过50万元的，应提交总经理办公会审批。"
  },
  "status": "active",
  "version": "v1.0"
}""")

add_heading(doc, "6. Agent 分工与编排", 1)
add_body(doc, "建议采用多 Agent 编排，而不是单一大模型节点。解析、抽取、标准化、冲突检测和审核协同需要拆开，便于追踪错误和替换模型。")
add_table(doc, ["Agent", "职责", "输入", "输出"], [
    ("文档解析 Agent", "解析文件结构，识别标题、表格、页码和条款编号", "原始文件", "结构化文本块"),
    ("文档分类 Agent", "判断资料类型、业务域、效力等级和抽取策略", "结构化文本", "分类标签"),
    ("知识整理 Agent", "生成摘要、标签、引用定位和知识切片", "文本块", "知识库条目"),
    ("规则抽取 Agent", "抽取义务、禁止、阈值、审批、必备条款和风险点", "知识切片", "候选规则"),
    ("规则标准化 Agent", "将自然语言候选规则转为统一 JSON/DSL", "候选规则", "规则草稿"),
    ("冲突检测 Agent", "比对现有规则，识别重复、冲突、过期和优先级问题", "规则草稿、规则库", "冲突报告"),
    ("审核协同 Agent", "为法务展示原文依据、影响范围和修改建议", "规则草稿、冲突报告", "审核任务"),
    ("规则发布 Agent", "发布审核通过规则，刷新检索索引和调用缓存", "已审核规则", "可调用规则包"),
], widths=[3.2, 5.4, 3.2, 3.1])

add_heading(doc, "7. 人工审核与版本治理", 1)
add_body(doc, "规则生效必须经过人工确认。系统应将审核界面设计成“候选规则 + 原文依据 + 适用范围 + 冲突提示 + 测试样例”的形式，而不是只展示 AI 总结文本。")
add_table(doc, ["状态", "含义", "允许动作"], [
    ("draft", "AI 生成的规则草稿，尚未进入审核", "编辑、删除、提交审核"),
    ("pending_review", "等待法务或管理员审核", "通过、驳回、退回修改"),
    ("active", "已生效，可在生成和审查中调用", "停用、创建新版本"),
    ("deprecated", "已废止，不再调用但保留历史", "查看、恢复为新版本"),
    ("conflicted", "与现有规则存在冲突，需处理", "合并、覆盖、保留高优先级规则"),
], widths=[3.0, 7.2, 4.6])
add_bullets(doc, [
    "版本号按规则维度维护，文档更新后不直接覆盖旧规则，而是生成新候选版本。",
    "每次规则调用都记录 rule_id、rule_version、合同 ID、命中条件和输出动作。",
    "法规和公司制度优先级高于历史合同和交易习惯；同等级规则按生效日期和人工优先级处理。",
])

add_heading(doc, "8. 在合同生成中的调用方式", 1)
add_body(doc, "合同生成不是简单调用知识库问答。推荐采用“生成前规则检索、生成中规则约束、生成后合规校验”的三段式调用。")
add_table(doc, ["时点", "调用内容", "系统行为"], [
    ("生成前", "合同类型、业务域、金额、主体、期限、交易模式", "召回模板、必填字段、审批规则、信息追问规则，先补齐关键信息"),
    ("生成中", "模板结构、用户输入、生效规则、标准条款", "按规则约束条款生成，不允许模型覆盖硬性规则"),
    ("生成后", "合同草稿、规则库、知识库依据", "逐条校验必备条款、禁止条款、审批条件和风险点，生成审查报告"),
    ("人工修改后", "用户修改痕迹、合同新版本", "再次运行差异审查，判断修改是否破坏规则合规性"),
], widths=[2.8, 5.4, 6.7])

add_heading(doc, "9. 在合同审查中的调用方式", 1)
add_body(doc, "审查流程应同时使用知识库和规则库。规则库负责确定性判断，知识库负责补充依据说明和相似条款参考。")
add_bullets(doc, [
    "必备条款审查：检查合同草稿是否包含适用规则要求的条款。",
    "禁止表达审查：识别不得承诺、不得免责、不得放弃权利等表达。",
    "审批触发审查：根据金额、期限、主体类型、付款方式判断是否需额外审批。",
    "依据标注：每个风险点必须展示来源文档、页码、章节和原文片段。",
    "修改建议：对可替换的条款给出标准条款建议；对不可自动修改的风险给出人工处理提示。",
])

add_heading(doc, "10. MVP 落地建议", 1)
add_table(doc, ["阶段", "建设内容", "验收标准"], [
    ("MVP 1", "支持 Word/PDF 上传、解析、切片、摘要、知识检索", "用户能按合同类型检索到原文依据，引用定位准确"),
    ("MVP 2", "支持必备条款、审批规则、风险提示三类规则抽取和审核", "AI 生成候选规则，人工审核后可启用"),
    ("MVP 3", "合同生成前调用规则，自动追问缺失字段", "生成采购/服务合同前能识别金额、期限、验收等缺失项"),
    ("MVP 4", "生成后审查并输出风险报告和依据标注", "每个风险点能关联规则 ID 与原文依据"),
    ("MVP 5", "规则版本、冲突检测、调用审计", "能查看规则变更历史和合同命中日志"),
], widths=[2.6, 7.2, 5.1])

add_heading(doc, "11. 关键页面建议", 1)
add_table(doc, ["页面", "核心能力"], [
    ("资料上传页", "批量上传、资料类型选择、业务域标注、保密等级、是否抽取规则"),
    ("知识库管理页", "原文预览、切片查看、摘要、标签、来源定位、重建索引"),
    ("候选规则页", "AI 候选规则列表、风险等级、来源引用、适用范围、置信度"),
    ("规则审核页", "规则编辑、原文对照、冲突报告、测试样例、通过/驳回"),
    ("规则库页", "按合同类型、规则类型、状态、版本、优先级检索和管理"),
    ("调用日志页", "查看合同生成和审查时命中的规则、模型输出、人工修改记录"),
], widths=[4.0, 10.9])

add_heading(doc, "12. 质量控制要求", 1)
add_bullets(doc, [
    "不得让 AI 抽取结果绕过人工审核直接生效，尤其是法规、审批、金额阈值和责任限制规则。",
    "每条规则必须绑定至少一个原文依据；没有依据的内容只能作为建议，不能作为生效规则。",
    "合同生成时必须区分硬性规则和建议性规则；硬性规则用于拦截或强制补充，建议性规则用于提示和推荐。",
    "知识库检索结果不能替代规则库判断；规则库判断也不能省略原文依据标注。",
    "系统必须保留规则版本、启停用记录、人工审核记录和合同调用日志，满足审计要求。",
])

add_heading(doc, "13. 一句话流程定义", 1)
add_body(doc, "用户上传资料后，系统先把资料解析并沉淀到可追溯知识库，再由 AI 从知识片段中提取候选规则，经过标准化、冲突检测和人工审核后形成可调用规则库；合同生成时先调用规则补齐信息和约束起草，合同审查时再调用规则和原文依据输出风险、修改建议和审批提醒。")

doc.save(OUT)
print(OUT)
